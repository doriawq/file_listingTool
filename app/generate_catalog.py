#!/usr/bin/env python3
import argparse
import os
import sys
from datetime import datetime
from pathlib import Path

def is_hidden(path: Path) -> bool:
    return any(part.startswith(".") for part in path.parts)


def iter_files(root: Path, include_hidden: bool, exclude: set[Path]):
    for dirpath, dirnames, filenames in os.walk(root):
        current = Path(dirpath)

        # Prune excluded directories
        pruned_dirnames = []
        for d in dirnames:
            dpath = current / d
            if dpath in exclude:
                continue
            if not include_hidden and is_hidden(dpath.relative_to(root)):
                continue
            pruned_dirnames.append(d)
        dirnames[:] = pruned_dirnames

        for name in filenames:
            fpath = current / name
            if fpath in exclude:
                continue
            rel = fpath.relative_to(root)
            if not include_hidden and is_hidden(rel):
                continue
            yield rel


def build_excel(rows, out_path: Path):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    wb = Workbook()
    ws = wb.active
    ws.title = "File Catalog"

    headers = ["No.", "Name (No Ext)"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="left")

    for idx, (name,) in enumerate(rows, start=1):
        ws.append([idx, name])

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 60

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def build_docx(rows, out_path: Path, root: Path, title: str | None = None):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("File Catalog", level=1)

    name = title.strip() if title else ""
    header_lines = []
    if name:
        header_lines.append(f"目录名称: {name}")
    header_lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    header_lines.append(f"Total files: {len(rows)}")
    doc.add_paragraph("\n".join(header_lines))

    doc.add_paragraph("Files:")
    if name:
        title_p = doc.add_paragraph()
        title_p.alignment = 1  # center
        title_run = title_p.add_run(name)
        title_run.font.name = "SimHei"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        title_run.font.size = Pt(16)

    for (name,) in rows:
        line = f"{name}"
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(line)
        run.font.name = "FangSong"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "FangSong")
        run.font.size = Pt(14)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def base_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent.parent


def parse_args():
    p = argparse.ArgumentParser(description="Generate file catalogs in Excel and Word formats.")
    default_root = str(base_dir() / "targets")
    p.add_argument(
        "--root",
        default=default_root,
        help="Root folder to scan (default: targets folder).",
    )
    p.add_argument(
        "--out-xlsx",
        default=str(Path(default_root) / "output/spreadsheet/file_catalog.xlsx"),
        help="Output Excel path.",
    )
    p.add_argument(
        "--out-docx",
        default=str(Path(default_root) / "output/doc/file_catalog.docx"),
        help="Output Word path.",
    )
    p.add_argument(
        "--include-hidden",
        action="store_true",
        help="Include hidden files/folders (starting with a dot).",
    )
    p.add_argument(
        "--exclude",
        action="append",
        default=[],
        help="Relative paths to exclude (can be repeated).",
    )
    p.add_argument(
        "--title",
        default="",
        help="Optional directory name shown in Word header.",
    )
    return p.parse_args()


def main():
    args = parse_args()
    run_catalog(
        root=args.root,
        out_xlsx=args.out_xlsx,
        out_docx=args.out_docx,
        include_hidden=args.include_hidden,
        exclude=args.exclude,
        title=args.title,
    )


def run_catalog(
    *,
    root: str,
    out_xlsx: str,
    out_docx: str,
    include_hidden: bool = False,
    exclude: list[str] | None = None,
    title: str = "",
):
    root_path = Path(root).resolve()

    exclude_set = {root_path / Path(p) for p in (exclude or [])}
    # Always exclude output and tmp by default unless user explicitly includes them
    exclude_set.update({root_path / "output", root_path / "tmp"})

    file_rows = []
    for rel in iter_files(root_path, include_hidden, exclude_set):
        file_rows.append((rel.stem,))

    file_rows.sort(key=lambda x: x[0].lower())

    build_excel(file_rows, Path(out_xlsx))
    build_docx(file_rows, Path(out_docx), root_path, title=title)

    print(f"Excel catalog: {Path(out_xlsx).resolve()}")
    print(f"Word catalog:  {Path(out_docx).resolve()}")
    print(f"Total files:   {len(file_rows)}")


if __name__ == "__main__":
    main()
