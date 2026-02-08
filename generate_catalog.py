#!/usr/bin/env python3
import argparse
import os
from datetime import datetime
from pathlib import Path

def is_hidden(path: Path) -> bool:
    return any(part.startswith(".") for part in path.parts)


def iter_files(root: Path, include_hidden: bool, exclude: set[Path]):
    for dirpath, dirnames, filenames in os.walk(root):
        current = Path(dirpath)

        # Prune excluded directories
        pruned_dirnames = []
        for d in dirnames:
            dpath = current / d
            if dpath in exclude:
                continue
            if not include_hidden and is_hidden(dpath.relative_to(root)):
                continue
            pruned_dirnames.append(d)
        dirnames[:] = pruned_dirnames

        for name in filenames:
            fpath = current / name
            if fpath in exclude:
                continue
            rel = fpath.relative_to(root)
            if not include_hidden and is_hidden(rel):
                continue
            yield rel


def build_excel(rows, out_path: Path):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    wb = Workbook()
    ws = wb.active
    ws.title = "File Catalog"

    headers = ["No.", "Name (No Ext)"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="left")

    for idx, (name,) in enumerate(rows, start=1):
        ws.append([idx, name])

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 60

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


def build_docx(rows, out_path: Path, root: Path, title: str | None = None):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("File Catalog", level=1)

    name = title.strip() if title else ""
    header_lines = []
    if name:
        header_lines.append(f"目录名称: {name}")
    header_lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    header_lines.append(f"Total files: {len(rows)}")
    doc.add_paragraph("\n".join(header_lines))

    doc.add_paragraph("Files:")
    if name:
        title_p = doc.add_paragraph()
        title_p.alignment = 1  # center
        title_run = title_p.add_run(name)
        title_run.font.name = "SimHei"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        title_run.font.size = Pt(16)

    for (name,) in rows:
        line = f"{name}"
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(line)
        run.font.name = "FangSong"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "FangSong")
        run.font.size = Pt(14)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def parse_args():
    p = argparse.ArgumentParser(description="Generate file catalogs in Excel and Word formats.")
    default_root = "/Users/Doria/Desktop/AI_projects/file_listingTool/targets"
    p.add_argument(
        "--root",
        default=default_root,
        help="Root folder to scan (default: targets folder).",
    )
    p.add_argument(
        "--out-xlsx",
        default=str(Path(default_root) / "output/spreadsheet/file_catalog.xlsx"),
        help="Output Excel path.",
    )
    p.add_argument(
        "--out-docx",
        default=str(Path(default_root) / "output/doc/file_catalog.docx"),
        help="Output Word path.",
    )
    p.add_argument(
        "--include-hidden",
        action="store_true",
        help="Include hidden files/folders (starting with a dot).",
    )
    p.add_argument(
        "--exclude",
        action="append",
        default=[],
        help="Relative paths to exclude (can be repeated).",
    )
    p.add_argument(
        "--title",
        default="",
        help="Optional directory name shown in Word header.",
    )
    return p.parse_args()


def main():
    args = parse_args()
    root = Path(args.root).resolve()

    exclude = {root / Path(p) for p in args.exclude}
    # Always exclude output and tmp by default unless user explicitly includes them
    exclude.update({root / "output", root / "tmp"})

    file_rows = []
    for rel in iter_files(root, args.include_hidden, exclude):
        file_rows.append((rel.stem,))

    file_rows.sort(key=lambda x: x[0].lower())

    build_excel(file_rows, Path(args.out_xlsx))
    build_docx(file_rows, Path(args.out_docx), root, title=args.title)

    print(f"Excel catalog: {Path(args.out_xlsx).resolve()}")
    print(f"Word catalog:  {Path(args.out_docx).resolve()}")
    print(f"Total files:   {len(file_rows)}")


if __name__ == "__main__":
    main()
